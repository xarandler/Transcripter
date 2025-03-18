import streamlit as st
import whisper
import os
from tempfile import NamedTemporaryFile
from docx import Document
from io import BytesIO

def seconds_to_srt_time(seconds):
    """Convert seconds to SRT time format"""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    millis = int((seconds - int(seconds)) * 1000)
    return f"{hours:02}:{minutes:02}:{secs:02},{millis:03}"

def generate_word_file(text):
    """Generate a Word document from the given text."""
    doc = Document()
    doc.add_heading('Transcription', level=1)
    doc.add_paragraph(text)
    
    # Save to BytesIO buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def main():
    st.set_page_config(
        page_title="Whisper Transcription",
        page_icon="🎙️",
        layout="wide",
        initial_sidebar_state="expanded"
    )

    st.title("🎙️ Whisper Audio Transcription")
    st.markdown("Optimized for transcription and translation into English or Swedish")

    device = "cpu"
    st.sidebar.markdown(f"**Device in use:** {device.upper()}")

    model_size = st.sidebar.selectbox(
        "Select Model Size",
        ("tiny", "base", "small", "medium", "large"),
        index=2,
        help="Larger models are more accurate but slower"
    )

    language_mode = st.sidebar.radio(
        "Language Mode",
        ("Auto-detect", "Force Swedish"),
        index=1,
        help="Force Swedish for unambiguous Swedish audio"
    )

    translation_mode = st.sidebar.radio(
        "Translation Mode",
        ("Transcribe only", "Translate to English", "Translate to Swedish"),
        index=0,
        help="Choose whether to translate the transcription"
    )

    @st.cache_resource
    def load_model(model_size):
        return whisper.load_model(model_size)

    with st.spinner(f"Loading {model_size} model..."):
        model = load_model(model_size)
        model = model.to(device)

    audio_file = st.file_uploader(
        "Upload audio file",
        type=["wav", "mp3", "m4a", "ogg", "flac"],
        help="Supports common audio formats"
    )

    if audio_file is not None:
        file_ext = os.path.splitext(audio_file.name)[1]
        with NamedTemporaryFile(suffix=file_ext, delete=False) as temp_file:
            temp_file.write(audio_file.getbuffer())
            temp_path = temp_file.name

        st.audio(temp_path)

        if st.button("Transcribe Audio"):
            with st.spinner("Transcribing..."):
                try:
                    task = "transcribe" if translation_mode == "Transcribe only" else "translate"
                    if translation_mode == "Translate to Swedish":
                        language_mode = "Force Swedish"
                    
                    result = model.transcribe(
                        temp_path,
                        language="sv" if language_mode == "Force Swedish" else None,
                        task=task
                    )
                    
                    detected_lang = result.get("language", "unknown").upper()
                    st.sidebar.markdown(f"**Detected language:** {detected_lang}")

                    text = result["text"]
                    segments = result["segments"]
                except Exception as e:
                    st.error(f"Transcription failed: {e}")
                    st.stop()
                finally:
                    try:
                        os.unlink(temp_path)
                    except Exception as e:
                        st.error(f"Error cleaning temporary file: {e}")

            st.success("Transcription Complete")
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("Transcription Text")
                st.text_area("Full Text", text, height=300)

            with col2:
                st.subheader("Segments")
                for segment in segments:
                    st.markdown(f"**{seconds_to_srt_time(segment['start'])} - {seconds_to_srt_time(segment['end'])}**")
                    st.write(segment['text'])
                    st.divider()

            srt_content = "\n".join(
                f"{i}\n{seconds_to_srt_time(seg['start'])} --> {seconds_to_srt_time(seg['end'])}\n{seg['text'].strip()}\n"
                for i, seg in enumerate(segments, 1)
            )

            st.download_button(
                label="Download TXT",
                data=text,
                file_name="transcription.txt",
                mime="text/plain"
            )
            
            st.download_button(
                label="Download SRT",
                data=srt_content,
                file_name="transcription.srt",
                mime="text/plain"
            )
            
            word_file = generate_word_file(text)
            st.download_button(
                label="Download Word File",
                data=word_file,
                file_name="transcription.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()
